from docx import Document


def read_word(path):
    print(f'Start read {path}')
    word_doc = Document(path)
    result = {}
    spam = None
    for table in word_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '\n' in cell.text:
                    cell_text = cell.text.replace(' ', '').replace('\n', ' ')
                else:
                    cell_text = cell.text
                if (cell_text.startswith('Блок') or cell_text.startswith('Ячейка') or cell_text.startswith('Субблок'))\
                        and len(cell_text) < 20:
                    result[cell_text] = []
                    spam = cell_text
                elif cell_text.startswith('ЫК') and spam in result:
                    result[spam].append(cell_text[:11])
                    result[spam].append('Э3')
                elif spam in result:
                    if 1 <= len(cell_text) <= 3:
                        if cell_text.isnumeric():
                            result[spam].append(int(cell_text))
                        else:
                            result[spam].append(cell_text)
    return result


# фунцкция возвращает кортеж последних литер изменения Э3 и ПЭ3
def last_liter(value, first_value=0, second_value=0):
    # print(f'Start search last liter: {value}')
    if value:
        if 'ПЭ3' in value:
            index = value.index('ПЭ3')
            first_value = value[index - 1]
            if 'ТУ' in value:
                index = value.index('ТУ')
                second_value = value[index - 1]
            else:
                second_value = value[-1]
        else:
            first_value = value[-1]

    return f'{first_value}, {second_value}'


# преобразования словаря в список для записи в Word
def dict_mod(data_dict, number_of_files):
    result = []
    number_of_items = len(data_dict)
    for key, values in data_dict.items():
        result.append(key)
        while len(values) < number_of_files + 1:
            values.append(['0, 0'])
        for liter in values:
            for item in liter:
                result.append(item)

    return result, number_of_items


def write_word(write_list, number_of_items, number_of_files):
    doc = Document('Output.docx')
    table = doc.add_table(rows=number_of_items + 1, cols=number_of_files + 2)
    # применяем стиль для таблицы
    table.style = 'Table Grid'
    nums = ['name', 'code', 3, 9, 13, 22, 26, 28, 35, 36, 43, 51, 54, 58]
    for col, num in enumerate(nums[:number_of_files + 2]):
        cell = table.cell(0, col)
        cell.text = str(num)
    counter = 0
    for row in range(1, number_of_items):
        for col in range(number_of_files + 2):
            cell = table.cell(row, col)
            try:
                cell.text = str(write_list[counter])
                counter += 1
            except IndexError:
                print('IndexError')
                doc.save('Output.docx')
    doc.save('Output.docx')


def main():
    nums = [3, 9, 13, 22, 26, 28, 35, 36, 43, 51, 54, 58]
    nums_tested = [3, 9]
    number_of_files = len(nums)
    result_liter_data = {}
    for i in nums:
        liter_data = read_word(f'Н1/Н1_{i}.docx')
        print('End reading')
        for key, value in liter_data.items():
            if key not in result_liter_data:
                if value[0].startswith('ЫК'):
                    result_liter_data[key] = [[value[0]]]
                else:
                    result_liter_data[key] = [['Unknown']]

            result_liter_data[key].append([last_liter(value)])
    write_list, number_of_items = dict_mod(result_liter_data, number_of_files)
    write_word(write_list, number_of_items, number_of_files)


if __name__ == '__main__':
    main()





